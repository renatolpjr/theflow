#!/usr/bin/env python3
"""
Script to generate Word technical manual for The Flow English Trainer
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import requests

# Image URLs
IMAGE_URLS = {
    'architecture': 'https://static.abacusaicdn.net/images/5454cf4a-9563-4601-b6e4-b97d314f2251.png'
}

# Download architecture image
print("Downloading architecture image...")
response = requests.get(IMAGE_URLS['architecture'])
arch_image = '/tmp/architecture.png'
with open(arch_image, 'wb') as f:
    f.write(response.content)

# Corporate colors (using docx RGBColor)
BLUE_PRIMARY = RGBColor(0, 51, 153)
RED_ACCENT = RGBColor(178, 34, 52)

print("\n=== Creating Word Technical Manual ===\n")

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Modify heading styles
h1_style = doc.styles['Heading 1']
h1_style.font.name = 'Calibri'
h1_style.font.size = Pt(20)
h1_style.font.bold = True
h1_style.font.color.rgb = BLUE_PRIMARY

h2_style = doc.styles['Heading 2']
h2_style.font.name = 'Calibri'
h2_style.font.size = Pt(16)
h2_style.font.bold = True
h2_style.font.color.rgb = BLUE_PRIMARY

# Create code style
try:
    code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(9)
except:
    code_style = doc.styles['Normal']

print("Creating cover page...")
# CAPA
title = doc.add_heading('THE FLOW ENGLISH TRAINER', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(28)
title.runs[0].font.color.rgb = BLUE_PRIMARY

subtitle = doc.add_paragraph('MANUAL TÉCNICO COMPLETO')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(18)
subtitle.runs[0].font.bold = True

subtitle2 = doc.add_paragraph('Documentação de Arquitetura, Configuração e Deploy')
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle2.runs[0].font.size = Pt(14)

doc.add_paragraph()
version = doc.add_paragraph('Versão: 1.0')
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
date = doc.add_paragraph('Data: Novembro 2024')
date.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

print("Creating table of contents...")
# ÍNDICE
doc.add_heading('ÍNDICE', 1)
doc.add_paragraph('Este documento contém 15 seções principais cobrindo todos os aspectos técnicos do sistema.')
doc.add_paragraph()

toc_items = [
    "1. INTRODUÇÃO",
    "2. ARQUITETURA DO SISTEMA",
    "3. ESTRUTURA DO PROJETO",
    "4. TECNOLOGIAS E BIBLIOTECAS",
    "5. BANCO DE DADOS",
    "6. APIS E INTEGRAÇÕES",
    "7. AUTENTICAÇÃO E SEGURANÇA",
    "8. CONFIGURAÇÃO E INSTALAÇÃO",
    "9. DEPLOY E HOSPEDAGEM",
    "10. MANUTENÇÃO E OPERAÇÃO",
    "11. DESENVOLVIMENTO E CUSTOMIZAÇÃO",
    "12. TESTES",
    "13. BOAS PRÁTICAS",
    "14. APÊNDICES",
    "15. SUPORTE E CONTATO"
]

for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

print("Creating Section 1: Introduction...")
# 1. INTRODUÇÃO
doc.add_heading('1. INTRODUÇÃO', 1)

doc.add_heading('1.1 Visão Geral do Sistema', 2)
doc.add_paragraph(
    'The Flow English Trainer é uma plataforma web completa de aprendizado de inglês americano, '
    'focada em exercícios de listening e speaking com feedback automatizado por Inteligência Artificial. '
    'O sistema utiliza tecnologias de ponta como Google Cloud Text-to-Speech, OpenAI TTS e ElevenLabs '
    'para gerar áudio ultra-realista, além de análise de fala com Abacus AI para fornecer feedback '
    'personalizado aos alunos.'
)

doc.add_heading('1.2 Objetivos e Escopo', 2)
objectives = doc.add_paragraph()
objectives.add_run('Objetivos principais:\n').bold = True
doc.add_paragraph('• Fornecer prática intensiva de listening e speaking', style='List Bullet')
doc.add_paragraph('• Oferecer feedback imediato e personalizado via IA', style='List Bullet')
doc.add_paragraph('• Gamificar o aprendizado para aumentar engajamento', style='List Bullet')
doc.add_paragraph('• Permitir gestão completa de conteúdo via painel admin', style='List Bullet')
doc.add_paragraph('• Escalar para milhares de usuários simultâneos', style='List Bullet')

doc.add_heading('1.3 Público-Alvo deste Manual', 2)
doc.add_paragraph(
    'Este manual é destinado a desenvolvedores, administradores de sistema, DevOps engineers e '
    'profissionais de TI responsáveis pela instalação, configuração, manutenção e customização '
    'da plataforma The Flow English Trainer.'
)

doc.add_page_break()

print("Creating Section 2: System Architecture...")
# 2. ARQUITETURA DO SISTEMA
doc.add_heading('2. ARQUITETURA DO SISTEMA', 1)

doc.add_heading('2.1 Visão Geral da Arquitetura', 2)
doc.add_paragraph(
    'The Flow utiliza uma arquitetura moderna de aplicação web full-stack baseada em Next.js 14 '
    'com App Router, combinando frontend React e backend API Routes em um único projeto. '
    'A arquitetura é cloud-native, utilizando serviços gerenciados para banco de dados (PostgreSQL), '
    'armazenamento de arquivos (AWS S3) e serviços de IA (Google Cloud, OpenAI, ElevenLabs).'
)

doc.add_heading('2.2 Stack Tecnológico Completo', 2)
doc.add_paragraph('Frontend:', style='List Bullet').runs[0].bold = True
doc.add_paragraph('  • Next.js 14 (App Router, React Server Components)', style='List Bullet 2')
doc.add_paragraph('  • React 18 com TypeScript 5.2', style='List Bullet 2')
doc.add_paragraph('  • Tailwind CSS 3.3 + Shadcn/ui (Radix UI)', style='List Bullet 2')
doc.add_paragraph('  • Framer Motion (animações)', style='List Bullet 2')

doc.add_paragraph('Backend:', style='List Bullet').runs[0].bold = True
doc.add_paragraph('  • Next.js API Routes (RESTful APIs)', style='List Bullet 2')
doc.add_paragraph('  • Prisma ORM 6.7 + PostgreSQL 14+', style='List Bullet 2')
doc.add_paragraph('  • NextAuth.js 4.24 (autenticação JWT)', style='List Bullet 2')

doc.add_paragraph('Serviços de IA:', style='List Bullet').runs[0].bold = True
doc.add_paragraph('  • Google Cloud Text-to-Speech API', style='List Bullet 2')
doc.add_paragraph('  • OpenAI Text-to-Speech API', style='List Bullet 2')
doc.add_paragraph('  • ElevenLabs API', style='List Bullet 2')
doc.add_paragraph('  • Abacus AI (análise de fala)', style='List Bullet 2')

doc.add_paragraph('Infraestrutura:', style='List Bullet').runs[0].bold = True
doc.add_paragraph('  • AWS S3 (armazenamento de áudio)', style='List Bullet 2')
doc.add_paragraph('  • PostgreSQL (banco de dados relacional)', style='List Bullet 2')
doc.add_paragraph('  • Vercel / AWS / DigitalOcean (hosting)', style='List Bullet 2')

doc.add_heading('2.3 Diagrama de Arquitetura', 2)
doc.add_paragraph('Diagrama simplificado da arquitetura:')
doc.add_picture(arch_image, width=Inches(6))

doc.add_heading('2.4 Padrões de Design Utilizados', 2)
doc.add_paragraph('• MVC (Model-View-Controller) adaptado para Next.js', style='List Bullet')
doc.add_paragraph('• Repository Pattern (Prisma como camada de acesso a dados)', style='List Bullet')
doc.add_paragraph('• API Gateway Pattern (Next.js API Routes)', style='List Bullet')
doc.add_paragraph('• Component-Based Architecture (React)', style='List Bullet')
doc.add_paragraph('• Server-Side Rendering (SSR) e Static Site Generation (SSG)', style='List Bullet')

doc.add_page_break()

print("Creating Section 3: Project Structure...")
# 3. ESTRUTURA DO PROJETO
doc.add_heading('3. ESTRUTURA DO PROJETO', 1)

doc.add_heading('3.1 Organização de Diretórios', 2)
structure = """the_flow_english_trainer/
└── nextjs_space/
    ├── app/                      # Next.js 14 App Router
    │   ├── (auth)/              # Rotas de autenticação
    │   │   ├── login/
    │   │   └── register/
    │   ├── admin/               # Painel administrativo
    │   │   ├── lessons/
    │   │   ├── exercises/
    │   │   ├── challenges/
    │   │   ├── videos/
    │   │   └── settings/
    │   ├── dashboard/           # Dashboard do aluno
    │   ├── lessons/             # Visualização de aulas
    │   ├── exercises/           # Exercícios (listening/speaking)
    │   ├── api/                 # API Routes
    │   │   ├── auth/
    │   │   ├── lessons/
    │   │   ├── listening-exercises/
    │   │   ├── speaking-exercises/
    │   │   ├── challenges/
    │   │   └── admin/
    │   ├── layout.tsx           # Layout raiz
    │   └── page.tsx             # Página inicial
    ├── components/              # Componentes React reutilizáveis
    │   ├── ui/                  # Componentes Shadcn/ui
    │   ├── admin/               # Componentes do admin
    │   └── exercises/           # Componentes de exercícios
    ├── lib/                     # Bibliotecas e utilitários
    │   ├── prisma.ts            # Cliente Prisma
    │   ├── auth.ts              # Configuração NextAuth
    │   └── utils.ts             # Funções utilitárias
    ├── prisma/
    │   ├── schema.prisma        # Schema do banco de dados
    │   └── seed.ts              # Dados iniciais
    ├── public/                  # Arquivos estáticos
    ├── .env                     # Variáveis de ambiente
    ├── next.config.js           # Configuração Next.js
    ├── package.json             # Dependências
    ├── tailwind.config.ts       # Configuração Tailwind
    └── tsconfig.json            # Configuração TypeScript
"""
code_para = doc.add_paragraph(structure)
code_para.style = 'Code'

doc.add_heading('3.2 Principais Componentes', 2)
doc.add_paragraph(
    'O sistema é organizado em componentes modulares e reutilizáveis. Os principais componentes incluem:'
)
doc.add_paragraph('• AudioPlayer: Reprodução de áudio com controles', style='List Bullet')
doc.add_paragraph('• ExerciseCard: Card de exercício com preview', style='List Bullet')
doc.add_paragraph('• AdminTable: Tabela de dados com CRUD', style='List Bullet')
doc.add_paragraph('• ProgressChart: Gráficos de progresso do aluno', style='List Bullet')
doc.add_paragraph('• VoiceSelector: Seletor de voz TTS', style='List Bullet')

doc.add_heading('3.3 Fluxo de Dados', 2)
doc.add_paragraph(
    '1. Cliente (Browser) faz requisição HTTP para Next.js\n'
    '2. Next.js roteia para API Route ou Server Component\n'
    '3. API Route valida autenticação (NextAuth)\n'
    '4. Prisma ORM consulta/modifica PostgreSQL\n'
    '5. Serviços externos (TTS, S3) são chamados quando necessário\n'
    '6. Resposta JSON é retornada ao cliente\n'
    '7. React atualiza UI com novos dados'
)

doc.add_page_break()

print("Creating Section 4: Technologies...")
# 4. TECNOLOGIAS E BIBLIOTECAS
doc.add_heading('4. TECNOLOGIAS E BIBLIOTECAS', 1)

doc.add_heading('4.1 Frontend', 2)
frontend_tech = [
    ('Next.js 14', 'Framework React com App Router, SSR, SSG e API Routes'),
    ('React 18', 'Biblioteca UI com Server Components e Suspense'),
    ('TypeScript 5.2', 'Superset JavaScript com tipagem estática'),
    ('Tailwind CSS 3.3', 'Framework CSS utility-first'),
    ('Shadcn/ui', 'Componentes acessíveis baseados em Radix UI'),
    ('Framer Motion', 'Biblioteca de animações'),
    ('React Hook Form', 'Gerenciamento de formulários'),
    ('Zod', 'Validação de schemas TypeScript-first'),
    ('Lucide Icons', 'Biblioteca de ícones')
]

for tech, desc in frontend_tech:
    p = doc.add_paragraph()
    p.add_run(f'{tech}: ').bold = True
    p.add_run(desc)

doc.add_heading('4.2 Backend', 2)
backend_tech = [
    ('Next.js API Routes', 'Endpoints RESTful serverless'),
    ('NextAuth.js 4.24', 'Autenticação completa com JWT'),
    ('Prisma ORM 6.7', 'ORM type-safe para PostgreSQL'),
    ('PostgreSQL 14+', 'Banco de dados relacional'),
    ('bcryptjs', 'Hash de senhas'),
    ('jsonwebtoken', 'Geração e validação de JWT')
]

for tech, desc in backend_tech:
    p = doc.add_paragraph()
    p.add_run(f'{tech}: ').bold = True
    p.add_run(desc)

doc.add_heading('4.3 Armazenamento e IA', 2)
ai_tech = [
    ('AWS SDK v3', 'Upload e gerenciamento de arquivos no S3'),
    ('Google Cloud TTS', 'Geração de áudio com vozes neurais'),
    ('OpenAI API', 'Text-to-Speech de alta qualidade'),
    ('ElevenLabs API', 'Vozes ultra-realistas'),
    ('Abacus AI', 'Análise e transcrição de fala')
]

for tech, desc in ai_tech:
    p = doc.add_paragraph()
    p.add_run(f'{tech}: ').bold = True
    p.add_run(desc)

doc.add_page_break()

print("Creating Section 5: Database...")
# 5. BANCO DE DADOS
doc.add_heading('5. BANCO DE DADOS', 1)

doc.add_heading('5.1 Schema Prisma', 2)
doc.add_paragraph(
    'O banco de dados utiliza PostgreSQL com Prisma ORM. O schema completo inclui mais de 20 modelos '
    'de dados cobrindo usuários, aulas, exercícios, desafios, gamificação e configurações.'
)

doc.add_heading('5.2 Principais Modelos de Dados', 2)

models_desc = [
    ('User', 'Usuários do sistema com autenticação e perfil'),
    ('Lesson', 'Aulas estruturadas com vocabulário e gramática'),
    ('ListeningExercise', 'Exercícios de compreensão auditiva'),
    ('SpeakingExercise', 'Exercícios de pronúncia e fala'),
    ('Challenge', 'Desafios gamificados'),
    ('VideoLesson', 'Vídeos educativos'),
    ('Badge', 'Conquistas e badges'),
    ('Achievement', 'Sistema de achievements'),
    ('ApiSettings', 'Configurações de serviços TTS'),
    ('RegistrationToken', 'Controle de cadastros via token')
]

for model, desc in models_desc:
    p = doc.add_paragraph()
    p.add_run(f'{model}: ').bold = True
    p.add_run(desc)

doc.add_heading('5.3 Relacionamentos', 2)
doc.add_paragraph('O schema utiliza relacionamentos complexos:')
doc.add_paragraph('• User 1:N Lesson (um usuário pode ter várias aulas)', style='List Bullet')
doc.add_paragraph('• Lesson 1:N ListeningExercise (uma aula tem vários exercícios)', style='List Bullet')
doc.add_paragraph('• User M:N Badge (usuários podem ter múltiplos badges)', style='List Bullet')
doc.add_paragraph('• Challenge M:N User (desafios podem ter múltiplos participantes)', style='List Bullet')

doc.add_heading('5.4 Migrations', 2)
doc.add_paragraph('Comandos Prisma para gerenciar migrations:')
migration_commands = """# Gerar cliente Prisma
npx prisma generate

# Criar migration
npx prisma migrate dev --name nome_da_migration

# Aplicar migrations em produção
npx prisma migrate deploy

# Resetar banco de dados (desenvolvimento)
npx prisma migrate reset

# Abrir Prisma Studio (GUI)
npx prisma studio
"""
code_para = doc.add_paragraph(migration_commands)
code_para.style = 'Code'

doc.add_page_break()

print("Creating Section 6: APIs and Integrations...")
# 6. APIS E INTEGRAÇÕES
doc.add_heading('6. APIS E INTEGRAÇÕES', 1)

doc.add_heading('6.1 API Routes do Next.js', 2)
doc.add_paragraph('Principais endpoints da aplicação:')

api_routes = [
    ('/api/auth/*', 'Autenticação (login, logout, session)'),
    ('/api/lessons/*', 'CRUD de aulas'),
    ('/api/listening-exercises/*', 'CRUD de exercícios de listening'),
    ('/api/speaking-exercises/*', 'CRUD de exercícios de speaking'),
    ('/api/challenges/*', 'CRUD de desafios'),
    ('/api/videos/*', 'CRUD de vídeos'),
    ('/api/admin/settings', 'Configurações de TTS e serviços'),
    ('/api/admin/users', 'Gerenciamento de usuários')
]

for route, desc in api_routes:
    p = doc.add_paragraph()
    p.add_run(f'{route}: ').bold = True
    p.add_run(desc)

doc.add_heading('6.2 Integração Google Cloud TTS', 2)
doc.add_paragraph('Configuração e uso do Google Cloud Text-to-Speech:')

google_tts_info = """Endpoint: https://texttospeech.googleapis.com/v1/text:synthesize
Método: POST
Autenticação: API Key (header: x-goog-api-key)

Exemplo de requisição:
{
  "input": {"text": "Hello, how are you?"},
  "voice": {
    "languageCode": "en-US",
    "name": "en-US-Neural2-A",
    "ssmlGender": "FEMALE"
  },
  "audioConfig": {
    "audioEncoding": "MP3",
    "speakingRate": 1.0,
    "pitch": 0.0
  }
}

Vozes disponíveis:
- en-US-Neural2-A (Female)
- en-US-Neural2-C (Female)
- en-US-Neural2-D (Male)
- en-US-Neural2-E (Female)
- en-US-Neural2-F (Female)
- en-US-Neural2-G (Female)
- en-US-Neural2-H (Female)
- en-US-Neural2-I (Male)
- en-US-Neural2-J (Male)
"""
code_para = doc.add_paragraph(google_tts_info)
code_para.style = 'Code'

doc.add_heading('6.3 Integração OpenAI TTS', 2)
openai_info = """Endpoint: https://api.openai.com/v1/audio/speech
Método: POST
Autenticação: Bearer Token

Exemplo de requisição:
{
  "model": "tts-1",
  "input": "Hello, how are you?",
  "voice": "alloy",
  "speed": 1.0
}

Vozes disponíveis: alloy, echo, fable, onyx, nova, shimmer
"""
code_para = doc.add_paragraph(openai_info)
code_para.style = 'Code'

doc.add_heading('6.4 Integração ElevenLabs', 2)
elevenlabs_info = """Endpoint: https://api.elevenlabs.io/v1/text-to-speech/{voice_id}
Método: POST
Autenticação: API Key (header: xi-api-key)

Exemplo de requisição:
{
  "text": "Hello, how are you?",
  "model_id": "eleven_monolingual_v1",
  "voice_settings": {
    "stability": 0.5,
    "similarity_boost": 0.5
  }
}
"""
code_para = doc.add_paragraph(elevenlabs_info)
code_para.style = 'Code'

doc.add_heading('6.5 AWS S3', 2)
doc.add_paragraph('Configuração para upload de arquivos de áudio:')
s3_info = """Bucket: theflow-audio-files
Região: us-east-1
SDK: @aws-sdk/client-s3

Exemplo de upload:
import { S3Client, PutObjectCommand } from '@aws-sdk/client-s3';

const s3Client = new S3Client({
  region: process.env.AWS_REGION,
  credentials: {
    accessKeyId: process.env.AWS_ACCESS_KEY_ID,
    secretAccessKey: process.env.AWS_SECRET_ACCESS_KEY
  }
});

const command = new PutObjectCommand({
  Bucket: 'theflow-audio-files',
  Key: `audio/${filename}`,
  Body: audioBuffer,
  ContentType: 'audio/mpeg'
});

await s3Client.send(command);
"""
code_para = doc.add_paragraph(s3_info)
code_para.style = 'Code'

doc.add_page_break()

print("Creating Section 7: Authentication and Security...")
# 7. AUTENTICAÇÃO E SEGURANÇA
doc.add_heading('7. AUTENTICAÇÃO E SEGURANÇA', 1)

doc.add_heading('7.1 NextAuth.js', 2)
doc.add_paragraph(
    'O sistema utiliza NextAuth.js para autenticação completa com estratégia JWT. '
    'Suporta login via credenciais (email/senha) com hash bcrypt.'
)

nextauth_config = """// lib/auth.ts
import NextAuth from 'next-auth';
import CredentialsProvider from 'next-auth/providers/credentials';
import { PrismaAdapter } from '@next-auth/prisma-adapter';
import prisma from './prisma';
import bcrypt from 'bcryptjs';

export const authOptions = {
  adapter: PrismaAdapter(prisma),
  providers: [
    CredentialsProvider({
      name: 'Credentials',
      credentials: {
        email: { label: "Email", type: "email" },
        password: { label: "Password", type: "password" }
      },
      async authorize(credentials) {
        const user = await prisma.user.findUnique({
          where: { email: credentials.email }
        });
        
        if (user && bcrypt.compareSync(credentials.password, user.password)) {
          return { id: user.id, email: user.email, role: user.role };
        }
        return null;
      }
    })
  ],
  session: { strategy: 'jwt' },
  callbacks: {
    async jwt({ token, user }) {
      if (user) token.role = user.role;
      return token;
    },
    async session({ session, token }) {
      session.user.role = token.role;
      return session;
    }
  }
};
"""
code_para = doc.add_paragraph(nextauth_config)
code_para.style = 'Code'

doc.add_heading('7.2 Sistema de Tokens de Registro', 2)
doc.add_paragraph(
    'O sistema utiliza tokens únicos para controlar cadastros. Apenas usuários com token válido '
    'podem se registrar, permitindo controle de acesso e convites.'
)

doc.add_heading('7.3 Controle de Acesso (RBAC)', 2)
doc.add_paragraph('Roles disponíveis:')
doc.add_paragraph('• admin: Acesso total ao painel administrativo', style='List Bullet')
doc.add_paragraph('• user: Acesso apenas às funcionalidades de aluno', style='List Bullet')

doc.add_heading('7.4 Segurança de API', 2)
doc.add_paragraph('Medidas de segurança implementadas:')
doc.add_paragraph('• Validação de sessão em todas as rotas protegidas', style='List Bullet')
doc.add_paragraph('• Sanitização de inputs com Zod', style='List Bullet')
doc.add_paragraph('• Hash de senhas com bcrypt (salt rounds: 10)', style='List Bullet')
doc.add_paragraph('• HTTPS obrigatório em produção', style='List Bullet')
doc.add_paragraph('• CORS configurado adequadamente', style='List Bullet')

doc.add_heading('7.5 Variáveis de Ambiente', 2)
doc.add_paragraph('Variáveis obrigatórias no arquivo .env:')
env_vars = """# Database
DATABASE_URL="postgresql://user:password@localhost:5432/theflow"

# NextAuth
NEXTAUTH_SECRET="seu-secret-aleatorio-aqui"
NEXTAUTH_URL="http://localhost:3000"

# AWS S3
AWS_ACCESS_KEY_ID="sua-access-key"
AWS_SECRET_ACCESS_KEY="sua-secret-key"
AWS_REGION="us-east-1"
AWS_BUCKET_NAME="theflow-audio-files"

# Google Cloud TTS
GOOGLE_TTS_API_KEY="sua-api-key"

# OpenAI
OPENAI_API_KEY="sk-..."

# ElevenLabs
ELEVENLABS_API_KEY="sua-api-key"
"""
code_para = doc.add_paragraph(env_vars)
code_para.style = 'Code'

doc.add_page_break()

print("Creating Section 8: Configuration and Installation...")
# 8. CONFIGURAÇÃO E INSTALAÇÃO
doc.add_heading('8. CONFIGURAÇÃO E INSTALAÇÃO', 1)

doc.add_heading('8.1 Pré-requisitos', 2)
doc.add_paragraph('Requisitos de sistema:')
doc.add_paragraph('• Node.js 18.0 ou superior', style='List Bullet')
doc.add_paragraph('• PostgreSQL 14 ou superior', style='List Bullet')
doc.add_paragraph('• Yarn 1.22 ou superior (ou npm 8+)', style='List Bullet')
doc.add_paragraph('• Conta AWS com acesso S3', style='List Bullet')
doc.add_paragraph('• API Key de pelo menos um serviço TTS (Google/OpenAI/ElevenLabs)', style='List Bullet')

doc.add_heading('8.2 Instalação Passo a Passo', 2)

installation_steps = """# Passo 1: Clonar o repositório
git clone <repo-url>
cd the_flow_english_trainer/nextjs_space

# Passo 2: Instalar dependências
yarn install

# Passo 3: Configurar variáveis de ambiente
cp .env.example .env
# Editar .env com suas credenciais

# Passo 4: Configurar banco de dados
yarn prisma generate
yarn prisma db push
yarn prisma db seed

# Passo 5: Executar em desenvolvimento
yarn dev

# Acesse: http://localhost:3000
"""
code_para = doc.add_paragraph(installation_steps)
code_para.style = 'Code'

doc.add_heading('8.3 Configuração de Serviços Externos', 2)

doc.add_paragraph('8.3.1 Google Cloud TTS', style='Heading 3')
doc.add_paragraph('1. Acesse Google Cloud Console (console.cloud.google.com)')
doc.add_paragraph('2. Crie um novo projeto ou selecione existente')
doc.add_paragraph('3. Ative a API "Cloud Text-to-Speech"')
doc.add_paragraph('4. Vá em "Credenciais" > "Criar credenciais" > "Chave de API"')
doc.add_paragraph('5. Copie a API Key e adicione ao .env como GOOGLE_TTS_API_KEY')
doc.add_paragraph('6. Configure a chave no painel admin em /admin/settings')

doc.add_paragraph('8.3.2 AWS S3', style='Heading 3')
aws_setup = """1. Acesse AWS Console (console.aws.amazon.com)
2. Vá em S3 e crie um novo bucket (ex: theflow-audio-files)
3. Configure permissões do bucket
4. Vá em IAM > Users > Create User
5. Anexe política "AmazonS3FullAccess"
6. Gere Access Key e Secret Key
7. Adicione credenciais ao .env
"""
doc.add_paragraph(aws_setup)

doc.add_paragraph('8.3.3 PostgreSQL', style='Heading 3')
postgres_setup = """# Instalar PostgreSQL (Ubuntu/Debian)
sudo apt update
sudo apt install postgresql postgresql-contrib

# Criar banco de dados
sudo -u postgres psql
CREATE DATABASE theflow;
CREATE USER theflow_user WITH PASSWORD 'senha_segura';
GRANT ALL PRIVILEGES ON DATABASE theflow TO theflow_user;

# Configurar DATABASE_URL no .env
DATABASE_URL="postgresql://theflow_user:senha_segura@localhost:5432/theflow"
"""
code_para = doc.add_paragraph(postgres_setup)
code_para.style = 'Code'

doc.add_page_break()

print("Creating Section 9: Deploy and Hosting...")
# 9. DEPLOY E HOSPEDAGEM
doc.add_heading('9. DEPLOY E HOSPEDAGEM', 1)

doc.add_heading('9.1 Preparação para Produção', 2)
doc.add_paragraph('Antes de fazer deploy:')
doc.add_paragraph('• Configurar todas as variáveis de ambiente de produção', style='List Bullet')
doc.add_paragraph('• Executar build de produção: yarn build', style='List Bullet')
doc.add_paragraph('• Testar build localmente: yarn start', style='List Bullet')
doc.add_paragraph('• Configurar domínio e SSL', style='List Bullet')
doc.add_paragraph('• Configurar backup de banco de dados', style='List Bullet')

doc.add_heading('9.2 Opção 1: Vercel (Recomendado)', 2)
doc.add_paragraph(
    'Vercel é a plataforma oficial do Next.js e oferece a melhor experiência de deploy.'
)

vercel_steps = """Passos:
1. Criar conta em vercel.com
2. Conectar repositório Git (GitHub/GitLab/Bitbucket)
3. Configurar projeto:
   - Framework Preset: Next.js
   - Build Command: yarn build
   - Output Directory: .next
4. Adicionar variáveis de ambiente no painel Vercel
5. Deploy automático a cada push

Vantagens:
✓ Zero configuração
✓ Edge Network global (CDN)
✓ Integração nativa com Next.js
✓ SSL automático
✓ CI/CD integrado
✓ Preview deployments para PRs

Limitações:
✗ Execução serverless (timeouts de 10s no free tier)
✗ Custos podem aumentar em escala
"""
doc.add_paragraph(vercel_steps)

doc.add_heading('9.3 Opção 2: AWS (EC2 + RDS)', 2)
doc.add_paragraph('Deploy completo em infraestrutura AWS com PM2 e Nginx.')

doc.add_heading('9.4 Opção 3: DigitalOcean App Platform', 2)
digitalocean_steps = """Passos:
1. Criar conta no DigitalOcean
2. Criar novo App
3. Conectar repositório Git
4. Configurar:
   - Build Command: yarn build
   - Run Command: yarn start
5. Adicionar PostgreSQL Database (Managed)
6. Configurar variáveis de ambiente
7. Deploy

Vantagens:
✓ Gerenciamento simplificado
✓ Escalabilidade automática
✓ Preço competitivo ($5-12/mês)
✓ Database gerenciado incluso
"""
doc.add_paragraph(digitalocean_steps)

doc.add_heading('9.5 Opção 4: Docker + Docker Compose', 2)
dockerfile = """# Dockerfile
FROM node:18-alpine
WORKDIR /app
COPY package.json yarn.lock ./
RUN yarn install --frozen-lockfile
COPY . .
RUN yarn prisma generate
RUN yarn build
EXPOSE 3000
CMD ["yarn", "start"]

# docker-compose.yml
version: '3.8'
services:
  app:
    build: .
    ports:
      - "3000:3000"
    environment:
      - DATABASE_URL=postgresql://user:pass@db:5432/theflow
    depends_on:
      - db
  db:
    image: postgres:14-alpine
    environment:
      POSTGRES_DB: theflow
      POSTGRES_USER: user
      POSTGRES_PASSWORD: pass
    volumes:
      - pgdata:/var/lib/postgresql/data
volumes:
  pgdata:

# Comandos:
docker-compose up -d
"""
code_para = doc.add_paragraph(dockerfile)
code_para.style = 'Code'

doc.add_heading('9.6 Opção 5: Railway', 2)
railway_steps = """Railway oferece deploy instantâneo com database incluso:

Passos:
1. Criar conta em railway.app
2. New Project > Deploy from GitHub
3. Selecionar repositório
4. Adicionar PostgreSQL database (automático)
5. Configurar variáveis de ambiente
6. Deploy automático

Vantagens:
✓ Setup instantâneo (< 5 minutos)
✓ Database PostgreSQL incluso
✓ Free tier generoso ($5 crédito/mês)
✓ SSL automático
"""
doc.add_paragraph(railway_steps)

doc.add_page_break()

print("Creating Section 10: Maintenance...")
# 10. MANUTENÇÃO E OPERAÇÃO
doc.add_heading('10. MANUTENÇÃO E OPERAÇÃO', 1)

doc.add_heading('10.1 Backup e Recuperação', 2)
backup_commands = """# Backup de banco de dados PostgreSQL
pg_dump -h host -U user -d theflow > backup_$(date +%Y%m%d_%H%M%S).sql

# Restaurar backup
psql -h host -U user -d theflow < backup_20241113_120000.sql

# Backup automatizado com cron (diário às 2h)
0 2 * * * /usr/bin/pg_dump -h localhost -U theflow_user theflow > /backups/theflow_$(date +\\%Y\\%m\\%d).sql

# Backup de arquivos S3
aws s3 sync s3://theflow-audio-files /backups/s3/
"""
code_para = doc.add_paragraph(backup_commands)
code_para.style = 'Code'

doc.add_heading('10.2 Atualizações de Dependências', 2)
update_commands = """# Verificar dependências desatualizadas
yarn outdated

# Atualizar dependências interativamente
yarn upgrade-interactive --latest

# Atualizar Prisma
yarn add -D prisma@latest
yarn add @prisma/client@latest
yarn prisma generate

# Atualizar Next.js
yarn add next@latest react@latest react-dom@latest
"""
code_para = doc.add_paragraph(update_commands)
code_para.style = 'Code'

doc.add_heading('10.3 Monitoramento de Performance', 2)
doc.add_paragraph('Métricas importantes a monitorar:')
doc.add_paragraph('• Tempo de resposta de APIs (< 200ms ideal)', style='List Bullet')
doc.add_paragraph('• Taxa de erros (< 1% ideal)', style='List Bullet')
doc.add_paragraph('• Uso de memória (< 80% ideal)', style='List Bullet')
doc.add_paragraph('• Uso de CPU (< 70% ideal)', style='List Bullet')
doc.add_paragraph('• Conexões de banco de dados (monitorar pool)', style='List Bullet')

doc.add_paragraph('Ferramentas recomendadas:')
doc.add_paragraph('• Sentry: Error tracking e performance monitoring', style='List Bullet')
doc.add_paragraph('• Vercel Analytics: Métricas de performance', style='List Bullet')
doc.add_paragraph('• CloudWatch: Monitoramento AWS', style='List Bullet')
doc.add_paragraph('• Datadog: Monitoramento completo', style='List Bullet')

doc.add_heading('10.4 Troubleshooting Comum', 2)

troubleshooting = """Problema: Erro de conexão com banco de dados
Solução:
- Verificar DATABASE_URL no .env
- Testar conexão: psql $DATABASE_URL
- Verificar se PostgreSQL está rodando
- Verificar firewall e security groups

Problema: Falha no upload de áudio para S3
Solução:
- Verificar credenciais AWS no .env
- Verificar permissões do bucket S3
- Verificar política IAM do usuário

Problema: TTS não está gerando áudio
Solução:
- Verificar API key no painel /admin/settings
- Verificar saldo/quota da API
- Verificar logs de erro no console

Problema: Erros de autenticação
Solução:
- Verificar NEXTAUTH_SECRET no .env
- Limpar cookies do navegador
- Verificar se sessão JWT está válida
"""
doc.add_paragraph(troubleshooting)

doc.add_page_break()

print("Creating Section 11: Development and Customization...")
# 11. DESENVOLVIMENTO E CUSTOMIZAÇÃO
doc.add_heading('11. DESENVOLVIMENTO E CUSTOMIZAÇÃO', 1)

doc.add_heading('11.1 Ambiente de Desenvolvimento', 2)
doc.add_paragraph('Configuração recomendada do VSCode com extensions: ESLint, Prettier, Prisma, Tailwind CSS IntelliSense.')

doc.add_heading('11.2 Adicionar Novos Exercícios', 2)
doc.add_paragraph('Passos para criar um novo tipo de exercício:')
doc.add_paragraph('1. Atualizar Schema Prisma e executar migration', style='List Bullet')
doc.add_paragraph('2. Criar API Routes para CRUD', style='List Bullet')
doc.add_paragraph('3. Criar Componentes React', style='List Bullet')
doc.add_paragraph('4. Atualizar Seed com dados de exemplo', style='List Bullet')

doc.add_heading('11.3 Customização de UI', 2)
ui_customization = """# Modificar cores (app/globals.css)
:root {
  --primary: 220 90% 56%;      /* Azul primário */
  --secondary: 355 78% 42%;    /* Vermelho secundário */
}

# Adicionar novos componentes Shadcn
npx shadcn-ui@latest add button
npx shadcn-ui@latest add dialog

# Personalizar logo
- Substituir arquivos em public/logo.png
"""
code_para = doc.add_paragraph(ui_customization)
code_para.style = 'Code'

doc.add_heading('11.4 Adicionar Novos Idiomas', 2)
doc.add_paragraph('Para expandir para outros idiomas:')
doc.add_paragraph('1. Adicionar campo "language" aos modelos', style='List Bullet')
doc.add_paragraph('2. Configurar vozes TTS para o novo idioma', style='List Bullet')
doc.add_paragraph('3. Criar conteúdo específico do idioma', style='List Bullet')
doc.add_paragraph('4. Implementar seletor de idioma na UI', style='List Bullet')

doc.add_page_break()

print("Creating Section 12: Testing...")
# 12. TESTES
doc.add_heading('12. TESTES', 1)

doc.add_heading('12.1 Estratégia de Testes', 2)
doc.add_paragraph('Recomenda-se implementar testes em três níveis: unitários, integração e end-to-end.')

doc.add_heading('12.2 Testes Unitários', 2)
unit_test_example = """# Instalar Jest e Testing Library
yarn add -D jest @testing-library/react @testing-library/jest-dom

# Exemplo de teste unitário
import { render, screen } from '@testing-library/react';
import { Button } from '@/components/ui/button';

describe('Button', () => {
  it('renders button with text', () => {
    render(<Button>Click me</Button>);
    expect(screen.getByText('Click me')).toBeInTheDocument();
  });
});

# Executar testes
yarn test
"""
code_para = doc.add_paragraph(unit_test_example)
code_para.style = 'Code'

doc.add_heading('12.3 Testes de Integração', 2)
doc.add_paragraph('Testar APIs com ferramentas como Postman ou Insomnia.')

doc.add_heading('12.4 Testes E2E', 2)
e2e_example = """# Instalar Playwright
yarn add -D @playwright/test

# Exemplo de teste E2E
import { test, expect } from '@playwright/test';

test('user can login', async ({ page }) => {
  await page.goto('http://localhost:3000/login');
  await page.fill('input[name="email"]', 'test@example.com');
  await page.fill('input[name="password"]', 'password123');
  await page.click('button[type="submit"]');
  await expect(page).toHaveURL('/dashboard');
});

# Executar testes E2E
yarn playwright test
"""
code_para = doc.add_paragraph(e2e_example)
code_para.style = 'Code'

doc.add_page_break()

print("Creating Section 13: Best Practices...")
# 13. BOAS PRÁTICAS
doc.add_heading('13. BOAS PRÁTICAS', 1)

doc.add_heading('13.1 Segurança', 2)
doc.add_paragraph('• NUNCA commitar arquivo .env para o repositório', style='List Bullet')
doc.add_paragraph('• Rotacionar API keys regularmente (a cada 90 dias)', style='List Bullet')
doc.add_paragraph('• Usar HTTPS em produção (obrigatório)', style='List Bullet')
doc.add_paragraph('• Validar e sanitizar todos os inputs do usuário', style='List Bullet')
doc.add_paragraph('• Implementar rate limiting em APIs públicas', style='List Bullet')
doc.add_paragraph('• Manter dependências atualizadas (security patches)', style='List Bullet')

doc.add_heading('13.2 Performance', 2)
doc.add_paragraph('• Otimizar imagens (usar next/image)', style='List Bullet')
doc.add_paragraph('• Implementar code splitting e lazy loading', style='List Bullet')
doc.add_paragraph('• Usar caching estratégico (Redis recomendado)', style='List Bullet')
doc.add_paragraph('• Criar índices apropriados no banco de dados', style='List Bullet')
doc.add_paragraph('• Minimizar chamadas a APIs externas', style='List Bullet')

doc.add_heading('13.3 Código', 2)
doc.add_paragraph('• Seguir convenções TypeScript e ESLint', style='List Bullet')
doc.add_paragraph('• Componentizar e reutilizar código', style='List Bullet')
doc.add_paragraph('• Escrever código autodocumentado', style='List Bullet')
doc.add_paragraph('• Adicionar comentários para lógica complexa', style='List Bullet')
doc.add_paragraph('• Usar tipos TypeScript rigorosos', style='List Bullet')

doc.add_page_break()

print("Creating Section 14: Appendices...")
# 14. APÊNDICES
doc.add_heading('14. APÊNDICES', 1)

doc.add_heading('14.1 Glossário de Termos', 2)
glossary = [
    ('API', 'Application Programming Interface - Interface de programação'),
    ('CDN', 'Content Delivery Network - Rede de distribuição de conteúdo'),
    ('JWT', 'JSON Web Token - Token de autenticação'),
    ('ORM', 'Object-Relational Mapping - Mapeamento objeto-relacional'),
    ('RBAC', 'Role-Based Access Control - Controle de acesso baseado em papéis'),
    ('SaaS', 'Software as a Service - Software como serviço'),
    ('SSR', 'Server-Side Rendering - Renderização no servidor'),
    ('TTS', 'Text-to-Speech - Conversão de texto em fala'),
    ('SSML', 'Speech Synthesis Markup Language - Linguagem de marcação para síntese de fala')
]

for term, definition in glossary:
    p = doc.add_paragraph()
    p.add_run(f'{term}: ').bold = True
    p.add_run(definition)

doc.add_heading('14.2 Referências e Links Úteis', 2)
references = [
    'Next.js Documentation: https://nextjs.org/docs',
    'Prisma Documentation: https://www.prisma.io/docs',
    'NextAuth.js Documentation: https://next-auth.js.org',
    'Google Cloud TTS: https://cloud.google.com/text-to-speech',
    'OpenAI API: https://platform.openai.com/docs',
    'ElevenLabs API: https://elevenlabs.io/docs',
    'AWS S3 Documentation: https://docs.aws.amazon.com/s3',
    'Tailwind CSS: https://tailwindcss.com/docs',
    'Shadcn/ui: https://ui.shadcn.com'
]

for ref in references:
    doc.add_paragraph(ref, style='List Bullet')

doc.add_heading('14.3 Changelog', 2)
changelog = """Versão 1.0 (Novembro 2024)
- Lançamento inicial
- Sistema completo de listening e speaking
- Integração com Google Cloud TTS, OpenAI e ElevenLabs
- Painel administrativo
- Sistema de gamificação
- Autenticação com NextAuth.js
"""
doc.add_paragraph(changelog)

doc.add_heading('14.4 Licenças e Créditos', 2)
doc.add_paragraph(
    'The Flow English Trainer utiliza diversas bibliotecas open-source. '
    'Consulte o arquivo package.json para lista completa de dependências e suas licenças.'
)

doc.add_page_break()

print("Creating Section 15: Support...")
# 15. SUPORTE E CONTATO
doc.add_heading('15. SUPORTE E CONTATO', 1)

doc.add_heading('15.1 Canais de Suporte', 2)
doc.add_paragraph('📧 Email: suporte@theflow.com', style='List Bullet')
doc.add_paragraph('💬 Discord: discord.gg/theflow', style='List Bullet')
doc.add_paragraph('📚 Documentação: docs.theflow.com', style='List Bullet')
doc.add_paragraph('🐛 GitHub Issues: github.com/theflow/issues', style='List Bullet')

doc.add_heading('15.2 Relatório de Bugs', 2)
doc.add_paragraph('Para reportar bugs, abra uma issue no GitHub incluindo:')
doc.add_paragraph('• Descrição detalhada do problema', style='List Bullet')
doc.add_paragraph('• Passos para reproduzir', style='List Bullet')
doc.add_paragraph('• Comportamento esperado vs. atual', style='List Bullet')
doc.add_paragraph('• Screenshots ou logs de erro', style='List Bullet')

doc.add_heading('15.3 Solicitação de Features', 2)
doc.add_paragraph(
    'Para solicitar novas funcionalidades, abra uma issue no GitHub com tag "feature request".'
)

doc.add_heading('15.4 Contribuições', 2)
doc.add_paragraph(
    'Contribuições são bem-vindas! Por favor, leia o guia de contribuição (CONTRIBUTING.md) '
    'antes de submeter pull requests.'
)

# Save Word document
doc_filename = '/home/ubuntu/TheFlow_Manual_Tecnico_Completo.docx'
doc.save(doc_filename)
print(f"\n✓ Word document created: {doc_filename}\n")

print("\n" + "="*60)
print("✓ WORD DOCUMENT CREATED SUCCESSFULLY!")
print("="*60)
print(f"\nFile: {doc_filename}")
print("\nFile is ready for download.\n")

